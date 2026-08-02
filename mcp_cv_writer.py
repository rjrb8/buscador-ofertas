import os
import re
from docx import Document
from mcp.server.fastmcp import FastMCP

mcp = FastMCP("cv-writer")


def normalizar_correo(correo: str) -> str:
    correo = correo.strip().lower()
    return re.sub(r"[^a-z0-9.\-_@]", "", correo)


@mcp.tool()
def generar_cv_docx(correo: str, contacto: dict, resumen_profesional: str,
                     experiencia: list, habilidades: list, educacion: list) -> str:
    """
    Genera un CV profesional en .docx en output/<correo_normalizado>_borrador2.docx.
    - contacto: {nombre, telefono, email, ubicacion, linkedin (opcional)}
    - experiencia: lista de {puesto, empresa, fechas, logros: [str]}
    - habilidades: lista de strings
    - educacion: lista de {titulo, institucion, fechas}
    Devuelve la ruta del archivo generado o un mensaje de ERROR.
    """
    try:
        os.makedirs("output", exist_ok=True)
        doc = Document()

        doc.add_heading(contacto.get("nombre", ""), level=0)

        contacto_linea = " | ".join(filter(None, [
            contacto.get("telefono"), contacto.get("email"),
            contacto.get("ubicacion"), contacto.get("linkedin")
        ]))
        p = doc.add_paragraph(contacto_linea)
        p.alignment = 1

        doc.add_heading("Resumen Profesional", level=1)
        doc.add_paragraph(resumen_profesional)

        doc.add_heading("Experiencia Laboral", level=1)
        for exp in experiencia:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('puesto','')} — {exp.get('empresa','')}").bold = True
            doc.add_paragraph(exp.get("fechas", ""))
            for logro in exp.get("logros", []):
                doc.add_paragraph(logro, style="List Bullet")

        doc.add_heading("Habilidades", level=1)
        doc.add_paragraph(" · ".join(habilidades))

        doc.add_heading("Educación", level=1)
        for edu in educacion:
            doc.add_paragraph(
                f"{edu.get('titulo','')} — {edu.get('institucion','')} ({edu.get('fechas','')})"
            )

        correo_normalizado = normalizar_correo(correo)
        ruta = f"output/{correo_normalizado}_borrador2.docx"
        doc.save(ruta)
        return ruta

    except Exception as e:
        return f"ERROR al generar el CV: {str(e)}"


if __name__ == "__main__":
    mcp.run()
