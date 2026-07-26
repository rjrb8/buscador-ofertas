import os
from docx import Document
from mcp.server.fastmcp import FastMCP

mcp = FastMCP("cv-reader")


@mcp.tool()
def read_cv_docx(ruta_archivo: str) -> str:
    """
    Lee un archivo .docx y devuelve su contenido como texto plano,
    incluyendo párrafos y tablas. Usar SIEMPRE que se necesite analizar
    un CV en formato Word.
    """
    if not os.path.exists(ruta_archivo):
        return f"ERROR: No existe el archivo en la ruta {ruta_archivo}"

    if not ruta_archivo.lower().endswith(".docx"):
        return "ERROR: El archivo no es .docx. Formato no soportado."

    try:
        doc = Document(ruta_archivo)
        contenido = []

        for parrafo in doc.paragraphs:
            if parrafo.text.strip():
                contenido.append(parrafo.text.strip())

        for tabla in doc.tables:
            for fila in tabla.rows:
                celdas = [celda.text.strip() for celda in fila.cells]
                if any(celdas):
                    contenido.append(" | ".join(celdas))

        texto_final = "\n".join(contenido)

        if not texto_final:
            return "ERROR: El documento está vacío o no se pudo extraer texto"

        return texto_final

    except Exception as e:
        return f"ERROR al leer el documento: {str(e)}"


if __name__ == "__main__":
    mcp.run()